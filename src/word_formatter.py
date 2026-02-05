"""
Word document formatter using python-docx.

Produces .docx files with proper formatting matching the VBA macro specifications.
"""

from pathlib import Path
from typing import Union

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .formatter import BaseFormatter, ExamType, BODY_FONT, BODY_SIZE, MARGIN_INCHES, CHOICE_INDENT_INCHES
from .parser import ParsedExam, Question, Answer


class WordFormatter(BaseFormatter):
    """
    Formatter that produces Microsoft Word (.docx) documents.

    Applies formatting rules:
    - Times New Roman, 12pt font
    - 0.5 inch margins
    - 2-line centered bold 20pt header (Title + section name)
    - Questions formatted as "#.\\t<text>"
    - Entire question block kept together on same page
    - Choices each on own line with blank line after each
    - Answer key on a new page
    - Practice exam answers: "#.\\t<letter>\\t<code ref>"
    - Final exam answers: "#.\\t<letter>)\\t<source> (<citation>)"
    """

    def get_extension(self) -> str:
        return '.docx'

    def format(self, exam: ParsedExam, output_path: Union[str, Path],
               exam_type: ExamType = ExamType.FINAL) -> Path:
        output_path = self.ensure_extension(output_path)

        doc = Document()
        self._setup_document(doc)

        title = exam.title or "Exam"

        # Questions header
        self._add_header(doc, title, "Questions")

        # Add questions section
        for question in exam.questions:
            self._add_question(doc, question)

        # Add answer key on a new page (if there are answers)
        if exam.answers:
            self._add_page_break(doc)
            self._add_header(doc, title, "Answer Key")
            for answer in exam.answers:
                self._add_answer(doc, answer, exam_type)

        doc.save(str(output_path))
        return output_path

    def _setup_document(self, doc: Document):
        """Configure document margins and default styles."""
        for section in doc.sections:
            section.top_margin = Inches(MARGIN_INCHES)
            section.bottom_margin = Inches(MARGIN_INCHES)
            section.left_margin = Inches(MARGIN_INCHES)
            section.right_margin = Inches(MARGIN_INCHES)

        style = doc.styles['Normal']
        font = style.font
        font.name = BODY_FONT
        font.size = Pt(BODY_SIZE)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT)

    def _add_header(self, doc: Document, title: str, section_name: str):
        """
        Add centered bold 20pt header with two lines:
        Line 1: Title (from input)
        Line 2: Section name (Questions or Answer Key)
        """
        # Line 1: Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_before = Pt(0)
        title_para.paragraph_format.space_after = Pt(0)
        run = title_para.add_run(title)
        run.font.name = BODY_FONT
        run.font.size = Pt(20)
        run.font.bold = True
        self._set_keep_with_next(title_para, True)

        # Line 2: Section name
        section_para = doc.add_paragraph()
        section_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        section_para.paragraph_format.space_before = Pt(0)
        section_para.paragraph_format.space_after = Pt(24)
        run = section_para.add_run(section_name)
        run.font.name = BODY_FONT
        run.font.size = Pt(20)
        run.font.bold = True

    def _add_question(self, doc: Document, question: Question):
        """
        Add a question with its choices to the document.

        Uses keep_with_next on all paragraphs except the last choice
        to ensure the entire question block stays on one page.
        """
        # Add question text
        q_para = doc.add_paragraph()
        q_para.add_run(f"{question.number}.\t{question.text}")
        self._format_paragraph(q_para, 'question')

        # Add choices - keep_with_next on all except the last
        num_choices = len(question.choices)
        for i, choice in enumerate(question.choices):
            c_para = doc.add_paragraph()
            c_para.add_run(f"{choice.letter}.\t{choice.text}")

            is_last = (i == num_choices - 1)
            if is_last:
                self._format_paragraph(c_para, 'choice_last')
            else:
                self._format_paragraph(c_para, 'choice')

            c_para.paragraph_format.left_indent = Inches(CHOICE_INDENT_INCHES)

    def _add_answer(self, doc: Document, answer: Answer, exam_type: ExamType):
        """Add an answer key entry to the document."""
        a_para = doc.add_paragraph()

        if exam_type == ExamType.PRACTICE:
            a_para.add_run(
                f"{answer.question_number}.\t{answer.answer_letter}\t{answer.payload}"
            )
        else:
            a_para.add_run(
                f"{answer.question_number}.\t{answer.answer_letter})\t{answer.payload}"
            )

        self._format_paragraph(a_para, 'answer')

    def _format_paragraph(self, para, para_type: str):
        """
        Apply formatting to a paragraph based on its type.

        Types:
        - 'question': Keep with next, blank line after
        - 'choice': Keep with next (not last choice), blank line after
        - 'choice_last': Last choice in a question block, blank line after
        - 'answer': Answer key entry, blank line after
        """
        fmt = para.paragraph_format

        # Common settings
        fmt.space_before = Pt(0)
        fmt.line_spacing = 1.0
        fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.left_indent = Pt(0)
        fmt.first_line_indent = Pt(0)

        for run in para.runs:
            run.font.name = BODY_FONT
            run.font.size = Pt(BODY_SIZE)
            run.font.bold = False

        if para_type == 'question':
            fmt.space_after = Pt(12)
            self._set_keep_together(para, True)
            self._set_keep_with_next(para, True)

        elif para_type == 'choice':
            fmt.space_after = Pt(12)
            self._set_keep_together(para, True)
            self._set_keep_with_next(para, True)  # Keep with next choice

        elif para_type == 'choice_last':
            fmt.space_after = Pt(12)
            self._set_keep_together(para, True)
            self._set_keep_with_next(para, False)  # Last choice, OK to break after

        elif para_type == 'answer':
            fmt.space_after = Pt(12)
            self._set_keep_together(para, True)
            self._set_keep_with_next(para, False)

    def _set_keep_together(self, para, value: bool):
        """Set paragraph keep-together property using XML."""
        pPr = para._element.get_or_add_pPr()
        keep_lines = pPr.find(qn('w:keepLines'))
        if value:
            if keep_lines is None:
                keep_lines = OxmlElement('w:keepLines')
                pPr.append(keep_lines)
        else:
            if keep_lines is not None:
                pPr.remove(keep_lines)

    def _set_keep_with_next(self, para, value: bool):
        """Set paragraph keep-with-next property using XML."""
        pPr = para._element.get_or_add_pPr()
        keep_next = pPr.find(qn('w:keepNext'))
        if value:
            if keep_next is None:
                keep_next = OxmlElement('w:keepNext')
                pPr.append(keep_next)
        else:
            if keep_next is not None:
                pPr.remove(keep_next)

    def _add_page_break(self, doc: Document):
        """Add a page break before the answer key."""
        doc.add_page_break()
